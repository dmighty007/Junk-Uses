#!/usr/bin/env python
# coding: utf-8

# In[2]:


from bs4 import BeautifulSoup
import docx
import urllib.request


# In[3]:


indx = "https://www.novelhall.com/Treasure-hunting-Into-the-Unknown-11886/"


# In[4]:


user_agent = 'Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.9.0.7) Gecko/2009021910 Firefox/3.0.7'
headers={'User-Agent':user_agent} 


# In[5]:


request = urllib.request.Request(indx,None,headers)
response = urllib.request.urlopen(request)


# In[6]:


soup = BeautifulSoup(response.read(),"html.parser")
link = soup.find_all(id = "post-11")
links = soup.find_all("a")
plinks = links[90:1250]


# In[7]:


def parse_html(elem):
    #elem = BeautifulSoup(html, features="html.parser")
    text = ''
    for e in elem:
        if isinstance(e, str):
            text += e.strip()
        elif e.name in ['br',  'p', 'h1', 'h2', 'h3', 'h4','tr', 'th']:
            text += '\n'
        elif e.name == 'li':
            text += '\n- '
    return text


# In[19]:


mydoc = docx.Document()
for link in plinks:
    string = link.get("href")
    test = "https://www.novelhall.com" + string
    request=urllib.request.Request(test,None,headers) #The assembled request
    response = urllib.request.urlopen(request)
    data = response.read() # The data u need
    soup = BeautifulSoup(data , 'html.parser')
    title = soup.find('title')
    T = title.prettify()[46:-21]
    mydoc.add_heading(T, level=0)
    txt = parse_html(soup.find(id="htmlContent"))
    mydoc.add_paragraph(txt)
    mydoc.add_page_break()
    print(T)
mydoc.save("my_written_file.docx")


# In[21]:


for link in plinks[-10:]:
    print(link)


# In[ ]:




